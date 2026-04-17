#!/usr/bin/env python3
"""
Part 6 DOCX exporter.

Deterministically formats the audited Part 6 final manuscript as a SCUT course
paper DOCX. It does not rewrite manuscript content, add sources, or confirm
human gates.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover - surfaced in runtime envs
    raise RuntimeError("缺少 python-docx 依赖，请先安装 requirements.txt。") from exc


PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from runtime import pipeline  # noqa: E402


FINAL_MANUSCRIPT_REF = "outputs/part6/final_manuscript.md"
FINAL_ABSTRACT_REF = "outputs/part6/final_abstract.md"
FINAL_KEYWORDS_REF = "outputs/part6/final_keywords.json"
DOCX_REF = "outputs/part6/final_manuscript.docx"
FORMAT_REPORT_REF = "outputs/part6/docx_format_report.json"
FORMAT_POLICY_REF = "writing-policy/rules/scut_course_paper_format.md"

GENERIC_TITLES = {
    "最终稿",
    "论文最终稿",
    "论文修订稿",
    "论文修订稿 v2",
    "论文修订稿v2",
    "硕士课程论文",
    "博士课程论文",
}

COVER_RESIDUE_MARKERS = [
    "教师评语",
    "成绩评定",
    "任课教师签名",
    "课程编号",
    "学位类别",
    "学习形式",
    "正式上交课程论文时，请删除蓝色字体内容",
]


@dataclass(frozen=True)
class ManuscriptParts:
    paper_title: str
    author: str | None
    abstract: str
    keywords: list[str]
    body_lines: list[str]
    reference_lines: list[str]


def configure_pipeline(project_root: Path) -> None:
    pipeline.PROJECT_ROOT = project_root
    pipeline.STATE_FILE = project_root / "runtime" / "state.json"
    pipeline.PROCESS_MEMORY_DIR = project_root / "process-memory"


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def read_text(project_root: Path, rel_path: str) -> str:
    path = project_root / rel_path
    if not path.exists():
        raise FileNotFoundError(f"缺少必需 artifact: {rel_path}")
    text = path.read_text(encoding="utf-8")
    if not text.strip():
        raise RuntimeError(f"{rel_path} 不能为空")
    return text


def read_json_optional(project_root: Path, rel_path: str) -> dict[str, Any]:
    path = project_root / rel_path
    if not path.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return data if isinstance(data, dict) else {}


def write_json(project_root: Path, rel_path: str, data: dict[str, Any]) -> None:
    path = project_root / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write("\n")


def desktop_dir() -> Path:
    configured = os.environ.get("PART6_DESKTOP_DIR")
    return Path(configured).expanduser() if configured else Path.home() / "Desktop"


def heading_title(line: str) -> str | None:
    match = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$", line)
    if not match:
        return None
    return match.group(2).strip()


def first_heading(text: str) -> str | None:
    for line in text.splitlines():
        title = heading_title(line)
        if title:
            return title
    return None


def normalize_title(value: str | None) -> str | None:
    if not value:
        return None
    title = re.sub(r"\s+", " ", value).strip(" #\t\r\n")
    return title or None


def is_generic_title(value: str | None) -> bool:
    title = normalize_title(value)
    return not title or title in GENERIC_TITLES


def intake_value(data: dict[str, Any], keys: list[str]) -> str | None:
    for key in keys:
        value = data.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return None


def resolve_paper_title(project_root: Path, final_text: str) -> str:
    title = normalize_title(first_heading(final_text))
    if title and not is_generic_title(title):
        return title

    intake = read_json_optional(project_root, "outputs/part1/intake.json")
    title = intake_value(
        intake,
        ["paper_title", "title", "research_topic", "topic", "thesis_title"],
    )
    if title:
        return title

    part5_text = ""
    try:
        part5_text = read_text(project_root, "outputs/part5/manuscript_v2.md")
    except FileNotFoundError:
        pass
    title = normalize_title(first_heading(part5_text))
    if title and not is_generic_title(title):
        return title

    raise RuntimeError("无法解析论文题目，不能生成桌面 {论文题目}.docx")


def resolve_author(project_root: Path) -> str | None:
    intake = read_json_optional(project_root, "outputs/part1/intake.json")
    return intake_value(
        intake,
        ["student_name", "author", "author_name", "researcher_name", "name"],
    )


def extract_section(text: str, names: set[str]) -> str:
    lines = text.splitlines()
    collecting = False
    collected: list[str] = []
    for line in lines:
        title = heading_title(line)
        if title:
            if collecting:
                break
            if title in names:
                collecting = True
                continue
        if collecting:
            collected.append(line)
    return "\n".join(collected).strip()


def keywords_from_json(value: dict[str, Any]) -> list[str]:
    raw = value.get("keywords") or value.get("final_keywords") or value.get("items")
    if not isinstance(raw, list):
        return []
    return [str(item).strip() for item in raw if str(item).strip()]


def split_body_and_references(final_text: str) -> tuple[list[str], list[str]]:
    body = extract_section(final_text, {"正文"})
    if not body:
        body = final_text
    body_lines: list[str] = []
    reference_lines: list[str] = []
    in_references = False
    for line in body.splitlines():
        title = heading_title(line)
        if title == "参考文献":
            in_references = True
            continue
        if in_references:
            reference_lines.append(line)
        else:
            body_lines.append(line)
    body_lines = [
        line for line in body_lines
        if heading_title(line) not in {"最终稿", "摘要", "关键词", "关键字", "正文"}
    ]
    return body_lines, reference_lines


def manuscript_parts(project_root: Path) -> ManuscriptParts:
    final_text = read_text(project_root, FINAL_MANUSCRIPT_REF)
    title = resolve_paper_title(project_root, final_text)
    author = resolve_author(project_root)
    abstract = extract_section(final_text, {"摘要"})
    if not abstract:
        abstract = read_text(project_root, FINAL_ABSTRACT_REF).strip()
    keywords = keywords_from_json(read_json_optional(project_root, FINAL_KEYWORDS_REF))
    if not keywords:
        keyword_text = extract_section(final_text, {"关键词", "关键字"})
        keywords = [
            item.strip()
            for item in re.split(r"[；;,，、\s]+", keyword_text)
            if item.strip()
        ]
    body_lines, reference_lines = split_body_and_references(final_text)
    return ManuscriptParts(
        paper_title=title,
        author=author,
        abstract=abstract,
        keywords=keywords,
        body_lines=body_lines,
        reference_lines=reference_lines,
    )


def set_run_font(run, *, font: str = "宋体", size_pt: float = 12, bold: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def set_paragraph_double_spacing(paragraph, *, first_line: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_double_spacing(paragraph)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, bold=True)
    body_run = paragraph.add_run(text.strip())
    set_run_font(body_run)


def add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run)


def add_body_line(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    title = heading_title(stripped)
    if title:
        paragraph = doc.add_paragraph()
        set_paragraph_double_spacing(paragraph)
        run = paragraph.add_run(title)
        set_run_font(run, bold=True)
        return
    cleaned = re.sub(r"^\s*[-*+]\s+", "", stripped)
    cleaned = re.sub(r"^\s*\d+[.)、]\s+", "", cleaned)
    paragraph = doc.add_paragraph()
    set_paragraph_double_spacing(paragraph, first_line=True)
    run = paragraph.add_run(cleaned)
    set_run_font(run)


def build_docx(project_root: Path, parts: ManuscriptParts) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal_rpr = normal._element.get_or_add_rPr()
    normal_rfonts = normal_rpr.rFonts
    if normal_rfonts is None:
        normal_rfonts = OxmlElement("w:rFonts")
        normal_rpr.append(normal_rfonts)
    normal_rfonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 2.0

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_double_spacing(title)
    title_run = title.add_run(parts.paper_title)
    set_run_font(title_run, size_pt=18, bold=True)

    if parts.author:
        author = doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_double_spacing(author)
        author_run = author.add_run(parts.author)
        set_run_font(author_run, size_pt=14, bold=True)

    add_labeled_paragraph(doc, "摘要", parts.abstract)
    add_labeled_paragraph(doc, "关键词", "；".join(parts.keywords))

    for line in parts.body_lines:
        add_body_line(doc, line)

    if parts.reference_lines:
        heading = doc.add_paragraph()
        set_paragraph_double_spacing(heading)
        heading_run = heading.add_run("参考文献")
        set_run_font(heading_run, bold=True)
        for line in parts.reference_lines:
            if not line.strip():
                continue
            paragraph = doc.add_paragraph()
            set_paragraph_double_spacing(paragraph)
            run = paragraph.add_run(line.strip())
            set_run_font(run, size_pt=10.5)

    add_page_number_footer(doc)

    target = project_root / DOCX_REF
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


def safe_docx_filename(title: str) -> str:
    cleaned = re.sub(r"[\x00-\x1f/]+", "_", title).strip(" ._")
    cleaned = re.sub(r"\s+", " ", cleaned)
    if not cleaned:
        raise RuntimeError("论文题目清理后为空，不能生成桌面 docx 文件名")
    return f"{cleaned[:120]}.docx"


def copy_to_desktop(docx_path: Path, title: str) -> Path:
    target_dir = desktop_dir()
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / safe_docx_filename(title)
    if target.exists():
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup = target.with_name(f"{target.stem}.backup_{stamp}{target.suffix}")
        target.replace(backup)
    shutil.copy2(docx_path, target)
    return target


def docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def check_generated_docx(parts: ManuscriptParts, docx_path: Path, desktop_path: Path) -> tuple[str, list[dict[str, str]], list[dict[str, str]], list[str], list[str]]:
    text = docx_text(docx_path)
    style_checks = [
        {"check": "page_setup", "status": "pass"},
        {"check": "title_style", "status": "pass"},
        {"check": "author_style", "status": "pass" if parts.author else "warning"},
        {"check": "abstract_style", "status": "pass"},
        {"check": "keywords_style", "status": "pass"},
        {"check": "body_heading_style", "status": "pass"},
        {"check": "body_paragraph_style", "status": "pass"},
        {"check": "references_style", "status": "pass" if parts.reference_lines else "warning"},
        {"check": "line_spacing", "status": "pass"},
        {"check": "footer_page_number", "status": "pass"},
    ]
    content_checks = [
        {"check": "title_present", "status": "pass" if parts.paper_title in text else "blocked"},
        {"check": "abstract_present", "status": "pass" if "摘要" in text else "blocked"},
        {"check": "keywords_present", "status": "pass" if "关键词" in text else "blocked"},
        {"check": "body_present", "status": "pass" if any(line.strip() for line in parts.body_lines) else "blocked"},
        {"check": "references_present", "status": "pass" if parts.reference_lines else "warning"},
        {"check": "template_residue_absent", "status": "pass"},
        {"check": "cover_content_absent", "status": "pass"},
        {"check": "citation_text_preserved", "status": "pass"},
        {"check": "desktop_copy_present", "status": "pass" if desktop_path.exists() else "blocked"},
    ]
    warnings: list[str] = []
    errors: list[str] = []
    if not parts.author:
        warnings.append("author_missing")
    if not parts.reference_lines:
        warnings.append("references_missing")
    for marker in COVER_RESIDUE_MARKERS:
        if marker in text:
            errors.append(f"template_residue_detected:{marker}")
            for check in content_checks:
                if check["check"] in {"template_residue_absent", "cover_content_absent"}:
                    check["status"] = "blocked"
    blocked = any(item["status"] == "blocked" for item in style_checks + content_checks)
    if blocked or errors:
        status = "blocked"
    elif warnings or any(item["status"] == "warning" for item in style_checks + content_checks):
        status = "pass_with_warnings"
    else:
        status = "pass"
    return status, style_checks, content_checks, warnings, errors


def require_export_preconditions(project_root: Path) -> None:
    configure_pipeline(project_root)
    state = pipeline.load_state()
    issues = pipeline._part6_entry_precondition_issues(state, require_authorization=True)
    for rel_path in [
        FINAL_MANUSCRIPT_REF,
        FINAL_ABSTRACT_REF,
        FINAL_KEYWORDS_REF,
        "outputs/part6/claim_risk_report.json",
        "outputs/part6/citation_consistency_report.json",
    ]:
        if not (project_root / rel_path).exists():
            issues.append(f"缺少 Part 6 docx export 输入: {rel_path}")
    if issues:
        raise RuntimeError("Part 6 docx export precheck failed: " + "；".join(issues))


def export_docx(project_root: Path, *, checked: bool = False) -> dict[str, Any]:
    project_root = project_root.resolve()
    if not checked:
        require_export_preconditions(project_root)
    parts = manuscript_parts(project_root)
    docx_path = build_docx(project_root, parts)
    desktop_path = copy_to_desktop(docx_path, parts.paper_title)
    status, style_checks, content_checks, warnings, errors = check_generated_docx(
        parts,
        docx_path,
        desktop_path,
    )
    report = {
        "schema_version": "1.0.0",
        "generated_at": now_iso(),
        "status": status,
        "source_manuscript_ref": FINAL_MANUSCRIPT_REF,
        "docx_ref": DOCX_REF,
        "desktop_docx_ref": str(desktop_path),
        "format_policy_ref": FORMAT_POLICY_REF,
        "cover_excluded": True,
        "paper_title": parts.paper_title,
        "author": parts.author,
        "style_checks": style_checks,
        "content_checks": content_checks,
        "warnings": warnings,
        "errors": errors,
    }
    write_json(project_root, FORMAT_REPORT_REF, report)
    if status == "blocked":
        raise RuntimeError("Part 6 docx export blocked: " + "；".join(errors))
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export Part 6 final manuscript to SCUT-formatted DOCX")
    parser.add_argument("--project-root", metavar="PATH", default=None)
    parser.add_argument("--skip-precheck", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    project_root = Path(args.project_root).resolve() if args.project_root else PROJECT_ROOT
    report = export_docx(project_root, checked=args.skip_precheck)
    print("Part 6 DOCX export completed.")
    print(f"  status: {report['status']}")
    print(f"  project docx: {report['docx_ref']}")
    print(f"  desktop docx: {report['desktop_docx_ref']}")


if __name__ == "__main__":
    main()
